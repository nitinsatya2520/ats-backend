from flask import Flask, request, jsonify
from flask_cors import CORS
from io import BytesIO
import PyPDF2
import pdfplumber
import docx
import spacy
import requests

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "https://kns-ats-resume-checker.vercel.app"}})


# Load spaCy model
nlp = spacy.load("en_core_web_sm")

GOOGLE_KG_API_URL = "https://kgsearch.googleapis.com/v1/entities:search"
GOOGLE_KG_API_KEY = "AIzaSyBPyu5zoMFhKUAcV5SHmV1-xhoyVFL3HGY"

# Function: Extract text from PDF
def extract_text_from_pdf(file):
    text = ""
    try:
        with pdfplumber.open(file) as pdf:
            text = "\n".join([page.extract_text() or "" for page in pdf.pages])
    except Exception as e:
        return "", f"Error extracting text: {str(e)}"

    return text.strip(), None


# Function: Extract text from DOCX
def extract_text_from_docx(file):
    text = ""
    try:
        doc = docx.Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return "", f"Error extracting text: {str(e)}"

    return text.strip(), None


# Function: Check formatting issues
def check_formatting(file, file_type):
    issues = []
    try:
        if file_type == "pdf":
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    if page.extract_tables():
                        issues.append("Resume contains tables, which may not be ATS-friendly.")
                    if page.images:
                        issues.append("Resume contains images, which may not be ATS-friendly.")
        elif file_type == "docx":
            doc = docx.Document(file)
            for table in doc.tables:
                issues.append("Resume contains tables, which may not be ATS-friendly.")
            for shape in doc.inline_shapes:
                issues.append("Resume contains images, which may not be ATS-friendly.")
    except Exception as e:
        issues.append(f"Error during formatting check: {str(e)}")

    return issues


def fetch_google_kg_skills(job_title):
    """Fetch related skills using Google Knowledge Graph API."""
    try:
        response = requests.get(
            GOOGLE_KG_API_URL, 
            params={"query": job_title, "key": GOOGLE_KG_API_KEY, "limit": 10}
        )
        data = response.json()
        skills = {item["name"].lower() for item in data.get("itemListElement", []) if "name" in item}
        return skills
    except Exception as e:
        print(f"Google KG API Error: {e}")
        return set()

# Function: Extract keywords using NLP
def extract_keywords(text):
    doc = nlp(text)
    return {token.lemma_.lower() for token in doc if token.is_alpha and not token.is_stop}


# Function: Categorize keywords
def categorize_keywords(resume_keywords):
    categories = {
        "skills": [
            "python", "flask", "django", "javascript", "typescript", "react", "angular", "vue", 
            "html", "css", "tailwind", "sql", "mysql", "postgresql", "mongodb", "firebase", 
            "git", "docker", "kubernetes", "linux", "bash", "aws", "azure", "gcp",
            "tensorflow", "pytorch", "machine learning", "deep learning", "data science",
            "nlp", "computer vision", "big data", "hadoop", "spark", "cloud computing",
            "blockchain", "cryptography", "cyber security", "ethical hacking", "penetration testing"
        ],
        "experience": [
            "developer", "software engineer", "internship", "freelance", "full-time", "part-time",
            "contract", "consulting", "project management", "team lead", "scrum master", 
            "agile development", "waterfall", "CI/CD", "devops", "system administration"
        ],
        "education": [
            "bachelor", "master", "mba", "phd", "university", "college", "diploma", 
            "bootcamp", "certification", "coursework", "degree", "online courses", "training"
        ],
        "soft_skills": [
            "teamwork", "leadership", "problem-solving", "communication", "critical thinking",
            "time management", "multitasking", "adaptability", "creativity", "collaboration",
            "emotional intelligence", "decision making", "negotiation"
        ],
        "tools": [
            "jira", "confluence", "notion", "figma", "adobe xd", "photoshop", "illustrator",
            "vs code", "intellij", "pycharm", "eclipse", "xcode", "android studio", 
            "unity", "unreal engine", "blender", "autocad"
        ],
        "frameworks": [
            "express", "nestjs", "fastapi", "spring", "hibernate", "ruby on rails", "laravel",
            "flutter", "react native", "electron", "next.js", "nuxt.js"
        ],
        "databases": [
            "mysql", "postgresql", "sqlite", "mongodb", "firebase", "dynamodb", "cassandra", "redis"
        ],
        "cloud": [
            "aws", "azure", "gcp", "heroku", "netlify", "vercel", "digital ocean", "cloudflare"
        ]
    }

    category_scores = {cat: 0 for cat in categories}

    for cat, keywords in categories.items():
        matches = [kw for kw in keywords if kw in resume_keywords]
        category_scores[cat] = (len(matches) / len(keywords)) * 100 if keywords else 0

    return category_scores


def calculate_score(matched_keywords, jd_keywords, formatting_issues, category_scores):
    # A simple scoring mechanism could be: 
    score = len(matched_keywords) * 10  # For simplicity, we give 10 points for each matched keyword.
    
    # Add bonus for formatting issues or category scores, etc.
    score -= len(formatting_issues) * 5  # Deduct for each formatting issue found.
    
    # Optionally, include category scores for extra weight.
    score += sum(category_scores.values()) * 0.5  # Adjust the multiplier as necessary.
    
    # Normalize the score if needed (e.g., scale it to 100).
    return min(score, 100)


def calculate_weighted_score(resume_keywords, job_keywords, weights):
    score = 0
    total_weight = 0
    
    for category, weight in weights.items():
        resume_category_keywords = resume_keywords.get(category, [])
        job_category_keywords = job_keywords.get(category, [])
        
        common_keywords = set(resume_category_keywords).intersection(set(job_category_keywords))
        category_score = len(common_keywords) * weight
        
        score += category_score
        total_weight += weight
    
    return score / total_weight if total_weight > 0 else 0




# Function: Generate feedback
def generate_feedback(matched_keywords, missing_keywords, formatting_issues, category_scores, jd_keywords):
    return {
        "overall_score": calculate_score(matched_keywords, jd_keywords, formatting_issues, category_scores),
        "category_scores": category_scores,
        "matched_keywords": list(matched_keywords),
        "missing_keywords": list(missing_keywords),
        "issues": formatting_issues
    }


# Endpoint: Analyze resume
@app.route('/analyze', methods=['POST'])
def analyze_resume():
    try:
        file = request.files.get('resume')
        job_description = request.form.get('job_description')

        if not file or not job_description:
            return jsonify({"error": "Resume file and job description are required"}), 400

        file_bytes = BytesIO(file.read())
        file_type = file.filename.split(".")[-1].lower()

        if file_type == "pdf":
            resume_text, text_error = extract_text_from_pdf(file_bytes)
        elif file_type == "docx":
            resume_text, text_error = extract_text_from_docx(file_bytes)
        else:
            return jsonify({"error": "Unsupported file format. Only PDF and DOCX are allowed."}), 400

        if text_error:
            return jsonify({"error": text_error}), 400

        formatting_issues = check_formatting(file_bytes, file_type)

        # Extract keywords from resume & job description
        resume_keywords = extract_keywords(resume_text)
        jd_keywords = extract_keywords(job_description)

        # Extract job title (first noun phrase)
        job_title = next((chunk.text for chunk in nlp(job_description).noun_chunks), None)

        # Fetch additional skills using Google KG API
        google_kg_skills = fetch_google_kg_skills(job_title) if job_title else set()

        # Combine job description skills with Google KG skills
        combined_jd_keywords = jd_keywords | google_kg_skills

        matched_keywords = resume_keywords & combined_jd_keywords
        missing_keywords = combined_jd_keywords - resume_keywords

        # Categorize keywords
        category_scores = categorize_keywords(resume_keywords)

        # Generate feedback
        feedback = generate_feedback(matched_keywords, missing_keywords, formatting_issues, category_scores, combined_jd_keywords)

        return jsonify(feedback)

    except Exception as e:
        return jsonify({"error": f"An error occurred: {str(e)}"}), 500


# Run the app
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)

